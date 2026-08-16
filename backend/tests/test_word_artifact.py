import tempfile
import unittest
from pathlib import Path

from docx import Document as WordDocument

from app.service.word_artifact import create_word_document, edit_word_document


class WordArtifactTests(unittest.TestCase):
    def test_creates_structured_word_document(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "generated.docx"
            create_word_document(
                output,
                "Nano 测试文档",
                "自动生成与验证",
                [
                    {"type": "heading", "level": 1, "text": "摘要"},
                    {"type": "paragraph", "text": "这是正文。"},
                    {"type": "bullets", "items": ["第一项", "第二项"]},
                    {
                        "type": "table",
                        "headers": ["项目", "状态"],
                        "rows": [["Registry", "完成"], ["Word", "完成"]],
                        "column_widths": [2, 1],
                    },
                ],
            )
            document = WordDocument(output)

            self.assertEqual(document.core_properties.title, "Nano 测试文档")
            self.assertIn("这是正文。", [item.text for item in document.paragraphs])
            self.assertEqual(len(document.tables), 1)
            self.assertEqual(document.tables[0].cell(2, 1).text, "完成")

    def test_edits_copy_without_modifying_source(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "source.docx"
            output = root / "edited.docx"
            create_word_document(
                source, "原始标题", "", [{"type": "paragraph", "text": "旧内容"}]
            )
            summary = edit_word_document(
                source,
                output,
                [
                    {"type": "replace_text", "old": "旧内容", "new": "新内容"},
                    {"type": "append_blocks", "blocks": [
                        {"type": "heading", "level": 1, "text": "附录"},
                        {"type": "numbered", "items": ["检查版式"]},
                    ]},
                ],
            )

            self.assertEqual(summary["replacements"], 1)
            self.assertIn("旧内容", [item.text for item in WordDocument(source).paragraphs])
            self.assertIn("新内容", [item.text for item in WordDocument(output).paragraphs])


if __name__ == "__main__":
    unittest.main()
