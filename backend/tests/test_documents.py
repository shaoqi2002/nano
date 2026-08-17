import unittest
import hashlib
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from uuid import UUID
from unittest.mock import AsyncMock, patch

from docx import Document as WordDocument
from fastapi import UploadFile

from app.service.document import (
    InvalidDocumentError,
    create_document,
    decode_text,
    preview_kind_for,
    safe_filename,
    word_text,
)


class AsyncTransaction:
    async def __aenter__(self):
        return self

    async def __aexit__(self, exc_type, exc, traceback):
        return False


class FakeSession:
    def begin(self):
        return AsyncTransaction()
from app.service.document_cache import DocumentCache, DocumentCacheError
from app.service.document_indexer import DocumentIndexRequests
from app.service.embedding import (
    EmbeddingConfigurationError,
    embedding_is_configured,
    resolve_embedding_base_url,
)


class DocumentServiceTests(unittest.TestCase):
    def test_browser_embedding_key_can_configure_indexing(self) -> None:
        self.assertTrue(embedding_is_configured("browser-key"))

    def test_accepts_bailian_workspace_base_url(self) -> None:
        base_url = (
            "https://workspace.ap-southeast-1.maas.aliyuncs.com/compatible-mode/v1/"
        )
        self.assertEqual(resolve_embedding_base_url(base_url), base_url.rstrip("/"))

    def test_rejects_non_bailian_base_url(self) -> None:
        with self.assertRaises(EmbeddingConfigurationError):
            resolve_embedding_base_url("http://127.0.0.1:8000/v1")

    def test_index_requests_deduplicate_document_keys(self) -> None:
        requests = DocumentIndexRequests()
        document_id = UUID("00000000-0000-0000-0000-000000000001")
        requests.submit(document_id, "first-key")
        requests.submit(document_id, "latest-key")

        self.assertEqual(requests.take(), (document_id, "latest-key", None))
        self.assertIsNone(requests.take())

    def test_index_requests_keep_base_url_in_memory(self) -> None:
        requests = DocumentIndexRequests()
        document_id = UUID("00000000-0000-0000-0000-000000000001")
        requests.submit(
            document_id,
            "browser-key",
            "https://dashscope-us.aliyuncs.com/compatible-mode/v1",
        )

        self.assertEqual(requests.take(), (
            document_id,
            "browser-key",
            "https://dashscope-us.aliyuncs.com/compatible-mode/v1",
        ))

    def test_safe_filename_removes_client_paths(self) -> None:
        self.assertEqual(safe_filename("C:\\fakepath\\报告.pdf"), "报告.pdf")
        self.assertEqual(safe_filename("../../notes.md"), "notes.md")

    def test_preview_kind_accepts_supported_extensions_case_insensitively(self) -> None:
        self.assertEqual(preview_kind_for("manual.PDF"), "pdf")
        self.assertEqual(preview_kind_for("notes.md"), "markdown")
        self.assertEqual(preview_kind_for("report.docx"), "word")

    def test_preview_kind_rejects_active_content(self) -> None:
        with self.assertRaisesRegex(InvalidDocumentError, "不支持"):
            preview_kind_for("page.html")
        with self.assertRaisesRegex(InvalidDocumentError, "不支持"):
            preview_kind_for("image.svg")

    def test_decode_text_supports_utf8_and_common_chinese_encoding(self) -> None:
        self.assertEqual(decode_text("你好".encode()), "你好")
        self.assertEqual(decode_text("你好".encode("gb18030")), "你好")

    def test_word_text_extracts_paragraphs_and_tables(self) -> None:
        document = WordDocument()
        document.add_paragraph("标题")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "名称"
        table.cell(0, 1).text = "Nano"
        stream = BytesIO()
        document.save(stream)

        extracted = word_text(stream.getvalue())

        self.assertIn("标题", extracted)
        self.assertIn("名称\tNano", extracted)


class DocumentUploadTests(unittest.IsolatedAsyncioTestCase):
    async def test_upload_returns_persisted_document(self) -> None:
        payload = "上传回归测试".encode()
        upload = UploadFile(filename="notes.txt", file=BytesIO(payload))
        persisted = AsyncMock(side_effect=lambda _session, document: document)

        with (
            TemporaryDirectory() as directory,
            patch(
                "app.service.document.document_cache.store_stream",
                return_value=Path(directory) / "cached.txt",
            ),
            patch("app.service.document.upload_path") as upload_path_mock,
            patch("app.service.document.add_document", persisted),
        ):
            document = await create_document(FakeSession(), upload)

        self.assertIsNotNone(document)
        self.assertEqual(document.original_name, "notes.txt")
        self.assertEqual(document.size_bytes, len(payload))
        self.assertEqual(document.checksum_sha256, hashlib.sha256(payload).hexdigest())
        upload_path_mock.assert_called_once()
        persisted.assert_awaited_once()



class DocumentCacheTests(unittest.TestCase):
    def test_downloads_once_then_uses_cached_file(self) -> None:
        payload = b"cached document"
        checksum = hashlib.sha256(payload).hexdigest()
        calls = []
        with TemporaryDirectory() as directory:
            cache = DocumentCache(Path(directory), maximum_bytes=1024)

            def download(object_key: str, destination: Path) -> None:
                calls.append(object_key)
                destination.write_bytes(payload)

            first = cache.ensure("documents/1/original.pdf", len(payload), checksum, download)
            second = cache.ensure("documents/1/original.pdf", len(payload), checksum, download)

            self.assertEqual(first, second)
            self.assertEqual(first.read_bytes(), payload)
            self.assertEqual(calls, ["documents/1/original.pdf"])

    def test_rejects_download_with_wrong_checksum(self) -> None:
        payload = b"expected"
        checksum = hashlib.sha256(payload).hexdigest()
        with TemporaryDirectory() as directory:
            cache = DocumentCache(Path(directory), maximum_bytes=1024)

            with self.assertRaisesRegex(DocumentCacheError, "校验失败"):
                cache.ensure(
                    "documents/1/original.pdf",
                    len(payload),
                    checksum,
                    lambda _key, destination: destination.write_bytes(b"tampered"),
                )

            self.assertFalse(cache.path_for("documents/1/original.pdf").exists())

    def test_evicts_oldest_file_when_capacity_is_reached(self) -> None:
        first_payload = b"first"
        second_payload = b"second"
        with TemporaryDirectory() as directory:
            cache = DocumentCache(Path(directory), maximum_bytes=len(second_payload))
            first_path = cache.store_stream(
                "documents/1/original.txt",
                BytesIO(first_payload),
                len(first_payload),
                hashlib.sha256(first_payload).hexdigest(),
            )
            second_path = cache.store_stream(
                "documents/2/original.txt",
                BytesIO(second_payload),
                len(second_payload),
                hashlib.sha256(second_payload).hexdigest(),
            )

            self.assertFalse(first_path.exists())
            self.assertEqual(second_path.read_bytes(), second_payload)


if __name__ == "__main__":
    unittest.main()
