import base64
import unittest
from datetime import datetime, timezone
from io import BytesIO
from types import SimpleNamespace

from docx import Document as WordDocument
from langchain_core.messages import HumanMessage
from pypdf import PdfWriter
from pydantic import ValidationError

from app.schema.conversation import ChatAttachment, SendMessageRequest
from app.service.document import InvalidDocumentError
from app.core.config import DEEPSEEK_MODEL, DEEPSEEK_VISION_MODEL
from app.service.conversation import (
    convert_messages,
    human_message_content,
    model_for_messages,
    prepare_chat_attachments,
)


class ChatAttachmentTests(unittest.TestCase):
    def test_allows_an_attachment_only_message(self) -> None:
        request = SendMessageRequest(
            attachments=[ChatAttachment(
                kind="text",
                name="notes.txt",
                media_type="text/plain",
                content="hello",
            )]
        )
        self.assertEqual(request.message, "")

    def test_rejects_empty_messages_and_invalid_image_data(self) -> None:
        with self.assertRaises(ValidationError):
            SendMessageRequest()
        with self.assertRaises(ValidationError):
            ChatAttachment(
                kind="image", name="bad.png", media_type="image/png", data="%%%"
            )

    def test_builds_text_and_image_content_blocks(self) -> None:
        image_data = base64.b64encode(b"\x89PNG\r\n\x1a\nsmall image").decode()
        blocks = human_message_content("分析", [
            {
                "kind": "text", "name": "a.txt", "media_type": "text/plain",
                "content": "attachment body", "data": None,
            },
            {
                "kind": "image", "name": "a.png", "media_type": "image/png",
                "content": None, "data": image_data,
            },
        ])
        self.assertEqual(blocks[0], {"type": "text", "text": "分析"})
        self.assertIn("attachment body", blocks[1]["text"])
        self.assertTrue(blocks[2]["image_url"]["url"].startswith("data:image/png;base64,"))

    def test_routes_image_history_to_the_vision_model(self) -> None:
        text_message = SimpleNamespace(attachments=[])
        image_message = SimpleNamespace(attachments=[{"kind": "image"}])
        self.assertEqual(model_for_messages([text_message]), DEEPSEEK_MODEL)
        self.assertEqual(
            model_for_messages([text_message, image_message]), DEEPSEEK_VISION_MODEL
        )

    def test_history_restores_attachments_as_multimodal_messages(self) -> None:
        message = SimpleNamespace(
            role="user",
            content="看看",
            attachments=[{
                "kind": "text", "name": "a.md", "media_type": "text/markdown",
                "content": "history", "data": None,
            }],
            created_at=datetime.now(timezone.utc),
        )
        converted = convert_messages([message])
        self.assertIsInstance(converted[-1], HumanMessage)
        self.assertIn("history", converted[-1].content[1]["text"])


class ChatDocumentAttachmentTests(unittest.IsolatedAsyncioTestCase):
    async def test_extracts_docx_into_chat_context(self) -> None:
        document = WordDocument()
        document.add_heading("聊天附件", level=1)
        document.add_paragraph("这是直接读取的 DOCX 内容。")
        stream = BytesIO()
        document.save(stream)
        attachment = ChatAttachment(
            kind="document",
            name="notes.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            data=base64.b64encode(stream.getvalue()).decode(),
        )

        prepared = await prepare_chat_attachments([attachment.model_dump()])

        self.assertIn("直接读取", prepared[0]["content"])
        blocks = human_message_content("总结", prepared)
        self.assertIn("直接读取", blocks[1]["text"])

    async def test_rejects_pdf_without_extractable_text(self) -> None:
        writer = PdfWriter()
        writer.add_blank_page(width=100, height=100)
        stream = BytesIO()
        writer.write(stream)
        attachment = {
            "kind": "document",
            "name": "scan.pdf",
            "media_type": "application/pdf",
            "data": base64.b64encode(stream.getvalue()).decode(),
            "content": None,
        }
        with self.assertRaises(InvalidDocumentError):
            await prepare_chat_attachments([attachment])


if __name__ == "__main__":
    unittest.main()
