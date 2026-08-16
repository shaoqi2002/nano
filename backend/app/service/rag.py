import hashlib
import re
from dataclasses import dataclass
from pathlib import Path
from uuid import UUID

from docx import Document as WordDocument
from pypdf import PdfReader
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import (
    RAG_CHUNK_OVERLAP,
    RAG_CHUNK_SIZE,
    RAG_MAX_SCORE_DROP,
    RAG_MIN_SIMILARITY,
    RAG_QUERY_MIN_SIMILARITY,
    RAG_TOP_K,
)
from app.model.document import Document, DocumentChunk
from app.service.embedding import embed_texts


PARSER_VERSION = "1"


@dataclass(frozen=True)
class TextBlock:
    text: str
    page_number: int | None = None
    section_title: str | None = None


@dataclass(frozen=True)
class ParsedChunk:
    content: str
    page_number: int | None
    section_title: str | None
    content_hash: str


class DocumentParsingError(ValueError):
    pass


def _clean_text(text: str) -> str:
    text = text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    return re.sub(r"[ \t]+", " ", text).strip()


def _decode_text(path: Path) -> str:
    content = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentParsingError("无法识别文档文本编码")


def _extract_blocks(document: Document, path: Path) -> list[TextBlock]:
    if document.preview_kind == "pdf":
        try:
            reader = PdfReader(str(path))
            return [
                TextBlock(_clean_text(page.extract_text() or ""), page_number=index)
                for index, page in enumerate(reader.pages, start=1)
            ]
        except Exception as error:
            raise DocumentParsingError("PDF 无法解析或已加密") from error

    if document.preview_kind == "word":
        try:
            word = WordDocument(str(path))
        except Exception as error:
            raise DocumentParsingError("DOCX 无法解析") from error
        blocks: list[TextBlock] = []
        heading: str | None = None
        for paragraph in word.paragraphs:
            text = _clean_text(paragraph.text)
            if not text:
                continue
            if paragraph.style and paragraph.style.name.lower().startswith("heading"):
                heading = text[:255]
            else:
                blocks.append(TextBlock(text, section_title=heading))
        for table in word.tables:
            rows = [" | ".join(_clean_text(cell.text) for cell in row.cells) for row in table.rows]
            blocks.append(TextBlock("\n".join(rows), section_title=heading))
        return blocks

    if document.preview_kind in {"text", "markdown"}:
        text = _decode_text(path)
        if document.preview_kind == "markdown":
            blocks: list[TextBlock] = []
            heading: str | None = None
            buffer: list[str] = []
            for line in text.splitlines():
                match = re.match(r"^#{1,6}\s+(.+)$", line.strip())
                if match:
                    if buffer:
                        blocks.append(TextBlock("\n".join(buffer), section_title=heading))
                        buffer = []
                    heading = match.group(1).strip()[:255]
                else:
                    buffer.append(line)
            if buffer:
                blocks.append(TextBlock("\n".join(buffer), section_title=heading))
            return blocks
        return [TextBlock(text)]

    raise DocumentParsingError("该文件类型暂不支持 RAG 索引")


def _split_text(text: str) -> list[str]:
    text = _clean_text(text)
    if not text:
        return []
    size = max(200, RAG_CHUNK_SIZE)
    overlap = min(max(0, RAG_CHUNK_OVERLAP), size // 2)
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + size, len(text))
        if end < len(text):
            search_start = start + int(size * 0.6)
            candidates = [text.rfind(mark, search_start, end) for mark in ("\n", "。", "！", "？", ". ")]
            boundary = max(candidates)
            if boundary > start:
                end = boundary + 1
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = max(start + 1, end - overlap)
    return chunks


def parse_document_chunks(document: Document, path: Path) -> list[ParsedChunk]:
    chunks: list[ParsedChunk] = []
    for block in _extract_blocks(document, path):
        for content in _split_text(block.text):
            chunks.append(
                ParsedChunk(
                    content=content,
                    page_number=block.page_number,
                    section_title=block.section_title,
                    content_hash=hashlib.sha256(content.encode("utf-8")).hexdigest(),
                )
            )
    if not chunks:
        if document.preview_kind == "pdf":
            raise DocumentParsingError("PDF 未提取到文字；扫描件需要 OCR，当前版本暂未启用")
        raise DocumentParsingError("文档中没有可索引的文字")
    return chunks


async def retrieve_sources(
    session: AsyncSession,
    query: str,
    embedding_api_key: str | None = None,
) -> list[dict]:
    ready_document = await session.scalar(
        select(Document.id).where(Document.index_status == "ready").limit(1)
    )
    if ready_document is None:
        return []
    query_vector = (await embed_texts([query], embedding_api_key))[0]
    distance = DocumentChunk.embedding.cosine_distance(query_vector)
    statement = (
        select(DocumentChunk, Document, distance.label("distance"))
        .join(Document, Document.id == DocumentChunk.document_id)
        .where(Document.index_status == "ready")
        .order_by(distance)
        .limit(max(1, RAG_TOP_K))
    )
    rows = (await session.execute(statement)).all()
    candidates: list[dict] = []
    for chunk, document, raw_distance in rows:
        similarity = 1.0 - float(raw_distance)
        candidates.append(
            {
                "document_id": str(document.id),
                "document_name": document.original_name,
                "chunk_id": chunk.id,
                "page_number": chunk.page_number,
                "section_title": chunk.section_title,
                "excerpt": chunk.content[:240],
                "similarity": round(similarity, 4),
                "content": chunk.content,
            }
        )
    return select_relevant_sources(candidates)


def select_relevant_sources(candidates: list[dict]) -> list[dict]:
    """Reject unrelated queries and weak tail matches from vector retrieval."""
    if not candidates:
        return []

    ordered = sorted(
        candidates,
        key=lambda source: float(source["similarity"]),
        reverse=True,
    )
    best_similarity = float(ordered[0]["similarity"])
    query_threshold = max(RAG_MIN_SIMILARITY, RAG_QUERY_MIN_SIMILARITY)
    if best_similarity < query_threshold:
        return []

    score_floor = max(
        RAG_MIN_SIMILARITY,
        best_similarity - max(0.0, RAG_MAX_SCORE_DROP),
    )
    return [
        source
        for source in ordered
        if float(source["similarity"]) >= score_floor
    ]


def build_rag_context(sources: list[dict]) -> str:
    parts = []
    for index, source in enumerate(sources, start=1):
        location = f"，第 {source['page_number']} 页" if source["page_number"] else ""
        parts.append(
            f"[来源 {index}] {source['document_name']}{location}\n{source['content']}"
        )
    return "\n\n".join(parts)


def public_sources(sources: list[dict]) -> list[dict]:
    return [{key: value for key, value in source.items() if key != "content"} for source in sources]
