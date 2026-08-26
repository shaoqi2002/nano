import hashlib
import logging
from pathlib import Path, PurePosixPath
from uuid import UUID, uuid4

from docx import Document as WordDocument
from fastapi import UploadFile
from sqlalchemy.ext.asyncio import AsyncSession
from starlette.concurrency import run_in_threadpool

from app.core.config import DOCUMENT_MAX_BYTES
from app.model.document import Document
from app.repository.document import (
    add_document,
    get_document,
    list_documents,
    remove_document,
)
from app.service.document_cache import DocumentCacheError, document_cache
from app.service.object_storage import (
    delete_object,
    download_object,
    upload_object,
    upload_path,
)


ALLOWED_EXTENSIONS = {
    ".pdf": ("pdf", "application/pdf"),
    ".md": ("markdown", "text/markdown"),
    ".markdown": ("markdown", "text/markdown"),
    ".txt": ("text", "text/plain"),
    ".csv": ("text", "text/csv"),
    ".json": ("text", "application/json"),
    ".log": ("text", "text/plain"),
    ".docx": (
        "word",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ),
    ".png": ("image", "image/png"),
    ".jpg": ("image", "image/jpeg"),
    ".jpeg": ("image", "image/jpeg"),
    ".gif": ("image", "image/gif"),
    ".webp": ("image", "image/webp"),
}
TEXT_PREVIEW_MAX_BYTES = 5 * 1024 * 1024
logger = logging.getLogger(__name__)


class DocumentNotFoundError(Exception):
    pass


class InvalidDocumentError(ValueError):
    pass


def safe_filename(filename: str | None) -> str:
    normalized = (filename or "").replace("\\", "/")
    name = PurePosixPath(normalized).name.strip()
    if not name or name in {".", ".."}:
        raise InvalidDocumentError("文件名不能为空")
    if len(name) > 255:
        raise InvalidDocumentError("文件名不能超过 255 个字符")
    return name


def preview_kind_for(filename: str) -> str:
    extension = PurePosixPath(filename).suffix.lower()
    try:
        return ALLOWED_EXTENSIONS[extension][0]
    except KeyError as error:
        supported = ", ".join(sorted(ALLOWED_EXTENSIONS))
        raise InvalidDocumentError(f"不支持该文件格式；支持：{supported}") from error


async def create_document(
    session: AsyncSession,
    upload: UploadFile,
) -> Document:
    original_name = safe_filename(upload.filename)
    preview_kind = preview_kind_for(original_name)
    extension = PurePosixPath(original_name).suffix.lower()
    document_id = uuid4()
    object_key = f"documents/{document_id}/original{extension}"
    content_type = ALLOWED_EXTENSIONS[extension][1]

    digest = hashlib.sha256()
    size = 0
    while chunk := await upload.read(1024 * 1024):
        size += len(chunk)
        if size > DOCUMENT_MAX_BYTES:
            raise InvalidDocumentError(
                f"文件不能超过 {DOCUMENT_MAX_BYTES // (1024 * 1024)} MiB"
            )
        digest.update(chunk)

    if size == 0:
        raise InvalidDocumentError("不能上传空文件")

    cached_path: Path | None = None
    try:
        await upload.seek(0)
        cached_path = await run_in_threadpool(
            document_cache.store_stream,
            object_key,
            upload.file,
            size,
            digest.hexdigest(),
        )
    except DocumentCacheError:
        logger.warning("Unable to populate document cache after upload", exc_info=True)
    try:
        if cached_path is not None:
            await run_in_threadpool(
                upload_path,
                object_key,
                cached_path,
                content_type,
            )
        else:
            await upload.seek(0)
            await run_in_threadpool(
                upload_object,
                object_key,
                upload.file,
                content_type,
            )
    except Exception:
        if cached_path is not None:
            try:
                await run_in_threadpool(document_cache.remove, object_key)
            except DocumentCacheError:
                logger.warning("Unable to remove failed upload cache", exc_info=True)
        raise
    document = Document(
        id=document_id,
        original_name=original_name,
        object_key=object_key,
        content_type=content_type,
        preview_kind=preview_kind,
        size_bytes=size,
        checksum_sha256=digest.hexdigest(),
    )

    try:
        async with session.begin():
            return await add_document(session, document)
    except Exception:
        await run_in_threadpool(delete_object, object_key)
        try:
            await run_in_threadpool(document_cache.remove, object_key)
        except DocumentCacheError:
            logger.warning("Unable to remove rolled-back document cache", exc_info=True)
        raise


async def create_document_from_path(
    session: AsyncSession,
    path: Path,
    original_name: str | None = None,
) -> Document:
    """Persist a generated local artifact through the same document-library path."""
    name = safe_filename(original_name or path.name)
    preview_kind = preview_kind_for(name)
    extension = PurePosixPath(name).suffix.lower()
    content_type = ALLOWED_EXTENSIONS[extension][1]
    size = path.stat().st_size
    if size == 0:
        raise InvalidDocumentError("不能保存空文件")
    if size > DOCUMENT_MAX_BYTES:
        raise InvalidDocumentError(
            f"文件不能超过 {DOCUMENT_MAX_BYTES // (1024 * 1024)} MiB"
        )
    checksum = hashlib.sha256(path.read_bytes()).hexdigest()
    document_id = uuid4()
    object_key = f"documents/{document_id}/original{extension}"
    cached_path: Path | None = None
    try:
        with path.open("rb") as stream:
            cached_path = await run_in_threadpool(
                document_cache.store_stream,
                object_key,
                stream,
                size,
                checksum,
            )
    except DocumentCacheError:
        logger.warning("Unable to cache generated document", exc_info=True)
    await run_in_threadpool(upload_path, object_key, cached_path or path, content_type)
    document = Document(
        id=document_id,
        original_name=name,
        object_key=object_key,
        content_type=content_type,
        preview_kind=preview_kind,
        size_bytes=size,
        checksum_sha256=checksum,
    )
    try:
        async with session.begin():
            return await add_document(session, document)
    except Exception:
        await run_in_threadpool(delete_object, object_key)
        if cached_path:
            try:
                await run_in_threadpool(document_cache.remove, object_key)
            except DocumentCacheError:
                logger.warning("Unable to remove generated document cache", exc_info=True)
        raise
async def get_documents(session: AsyncSession) -> list[Document]:
    return await list_documents(session)


async def require_document(
    session: AsyncSession,
    document_id: UUID,
) -> Document:
    document = await get_document(session, document_id)
    if document is None:
        raise DocumentNotFoundError
    return document


async def delete_document(session: AsyncSession, document_id: UUID) -> None:
    async with session.begin():
        document = await get_document(session, document_id)
        if document is None:
            raise DocumentNotFoundError
        await run_in_threadpool(delete_object, document.object_key)
        try:
            await run_in_threadpool(document_cache.remove, document.object_key)
        except DocumentCacheError:
            logger.warning("Unable to remove deleted document cache", exc_info=True)
        await remove_document(session, document)


async def reindex_document(session: AsyncSession, document_id: UUID) -> Document:
    async with session.begin():
        document = await get_document(session, document_id)
        if document is None:
            raise DocumentNotFoundError
        if document.preview_kind == "image":
            raise InvalidDocumentError("图片需要 OCR，当前版本暂未启用")
        document.index_status = "pending"
        document.index_error = None
        document.indexed_at = None
        await session.flush()
        await session.refresh(document)
        return document


def decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise InvalidDocumentError("无法识别文档文本编码")


def word_text(content: bytes) -> str:
    from io import BytesIO
    from zipfile import BadZipFile, ZipFile

    try:
        with ZipFile(BytesIO(content)) as archive:
            entries = archive.infolist()
            if len(entries) > 10_000 or sum(item.file_size for item in entries) > 50 * 1024 * 1024:
                raise InvalidDocumentError("DOCX 解压后内容过大")
            if "word/document.xml" not in {item.filename for item in entries}:
                raise InvalidDocumentError("DOCX 缺少正文内容")
        document = WordDocument(BytesIO(content))
    except InvalidDocumentError:
        raise
    except BadZipFile as error:
        raise InvalidDocumentError("DOCX 文件无法解析") from error
    except Exception as error:
        raise InvalidDocumentError("DOCX 文件无法解析") from error

    blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            blocks.append("\t".join(cell.text for cell in row.cells))
    return "\n\n".join(blocks)


def pdf_text(content: bytes) -> str:
    from io import BytesIO
    from pypdf import PdfReader

    try:
        reader = PdfReader(BytesIO(content))
        blocks = [page.extract_text() or "" for page in reader.pages]
    except Exception as error:
        raise InvalidDocumentError("PDF 文件无法解析") from error
    return "\n\n".join(block for block in blocks if block.strip())


def cached_document_path(document: Document) -> Path:
    return document_cache.ensure(
        document.object_key,
        document.size_bytes,
        document.checksum_sha256,
        download_object,
    )


def document_text(document: Document) -> str:
    if document.preview_kind not in {"text", "markdown", "word"}:
        raise InvalidDocumentError("该文件不支持文本预览")
    path = cached_document_path(document)
    content = document_cache.read_bytes(path, TEXT_PREVIEW_MAX_BYTES)
    if document.preview_kind == "word":
        return word_text(content)
    return decode_text(content)
