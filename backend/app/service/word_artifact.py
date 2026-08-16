import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Any

from docx import Document as WordDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader


BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
GRAY = RGBColor(0x66, 0x66, 0x66)
CONTENT_WIDTH_DXA = 9360


class WordArtifactError(RuntimeError):
    pass


def _font(run, size: float, *, bold: bool = False, color=None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _configure_document(document: WordDocument, title: str) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    if "Document Quote" not in document.styles:
        quote = document.styles.add_style("Document Quote", WD_STYLE_TYPE.PARAGRAPH)
        quote.base_style = normal
        quote.font.italic = True
        quote.font.color.rgb = GRAY
        quote.paragraph_format.left_indent = Inches(0.35)
        quote.paragraph_format.right_indent = Inches(0.35)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _font(header.add_run(title[:80]), 8.5, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _font(footer.add_run("Nano Agent  ·  "), 8.5, color=GRAY)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def _set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _configure_table(table, column_widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    width = tbl_pr.first_child_found_in("w:tblW")
    width.set(qn("w:w"), str(sum(column_widths)))
    width.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    tbl_pr.append(indent)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in column_widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(value))
        grid.append(col)
    for row_index, row in enumerate(table.rows):
        for index, cell in enumerate(row.cells):
            cell.width = Inches(column_widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(column_widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    _font(run, 10, bold=row_index == 0)
            if row_index == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "F2F4F7")
                cell._tc.get_or_add_tcPr().append(shading)


def _add_table(document: WordDocument, block: dict[str, Any]) -> None:
    headers = [str(value) for value in block.get("headers") or []]
    rows = [[str(value) for value in row] for row in block.get("rows") or []]
    columns = len(headers) or (len(rows[0]) if rows else 0)
    if columns == 0 or any(len(row) != columns for row in rows):
        raise WordArtifactError("表格列数不一致")
    table = document.add_table(rows=1 if headers else 0, cols=columns)
    table.style = "Table Grid"
    if headers:
        for index, value in enumerate(headers):
            table.rows[0].cells[index].text = value
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    requested = block.get("column_widths") or []
    if requested and len(requested) == columns:
        total = sum(float(value) for value in requested)
        widths = [round(CONTENT_WIDTH_DXA * float(value) / total) for value in requested]
        widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    else:
        widths = [CONTENT_WIDTH_DXA // columns] * columns
        widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    _configure_table(table, widths)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def append_blocks(document: WordDocument, blocks: list[dict[str, Any]]) -> None:
    for block in blocks:
        kind = str(block.get("type") or "paragraph")
        if kind == "heading":
            level = min(max(int(block.get("level", 1)), 1), 3)
            document.add_heading(str(block.get("text") or ""), level=level)
        elif kind == "paragraph":
            document.add_paragraph(str(block.get("text") or ""))
        elif kind == "quote":
            document.add_paragraph(str(block.get("text") or ""), style="Document Quote")
        elif kind in {"bullets", "numbered"}:
            style = "List Bullet" if kind == "bullets" else "List Number"
            for item in block.get("items") or []:
                paragraph = document.add_paragraph(str(item), style=style)
                paragraph.paragraph_format.space_after = Pt(8)
                paragraph.paragraph_format.line_spacing = 1.167
        elif kind == "table":
            _add_table(document, block)
        elif kind == "page_break":
            document.add_page_break()
        else:
            raise WordArtifactError(f"不支持的内容块：{kind}")


def create_word_document(
    output_path: Path,
    title: str,
    subtitle: str,
    blocks: list[dict[str, Any]],
) -> None:
    document = WordDocument()
    _configure_document(document, title)
    title_paragraph = document.add_paragraph()
    title_paragraph.paragraph_format.space_after = Pt(4)
    _font(title_paragraph.add_run(title), 24, bold=True)
    if subtitle:
        subtitle_paragraph = document.add_paragraph()
        subtitle_paragraph.paragraph_format.space_after = Pt(16)
        _font(subtitle_paragraph.add_run(subtitle), 13, color=GRAY)
    append_blocks(document, blocks)
    document.core_properties.title = title
    document.core_properties.author = "Nano Agent"
    document.save(output_path)


def _replace_in_paragraph(paragraph, old: str, new: str) -> int:
    if old not in paragraph.text:
        return 0
    if any(old in run.text for run in paragraph.runs):
        count = 0
        for run in paragraph.runs:
            occurrences = run.text.count(old)
            run.text = run.text.replace(old, new)
            count += occurrences
        return count
    count = paragraph.text.count(old)
    paragraph.text = paragraph.text.replace(old, new)
    return count


def edit_word_document(
    source_path: Path,
    output_path: Path,
    operations: list[dict[str, Any]],
) -> dict[str, int]:
    document = WordDocument(source_path)
    counts = {"replacements": 0, "deleted_paragraphs": 0, "appended_blocks": 0}
    for operation in operations:
        kind = str(operation.get("type") or "")
        if kind == "replace_text":
            old = str(operation.get("old") or "")
            if not old:
                raise WordArtifactError("replace_text.old 不能为空")
            new = str(operation.get("new") or "")
            for paragraph in document.paragraphs:
                counts["replacements"] += _replace_in_paragraph(paragraph, old, new)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            counts["replacements"] += _replace_in_paragraph(paragraph, old, new)
        elif kind == "delete_paragraph":
            contains = str(operation.get("contains") or "")
            for paragraph in list(document.paragraphs):
                if contains and contains in paragraph.text:
                    paragraph._element.getparent().remove(paragraph._element)
                    counts["deleted_paragraphs"] += 1
        elif kind == "append_blocks":
            blocks = list(operation.get("blocks") or [])
            append_blocks(document, blocks)
            counts["appended_blocks"] += len(blocks)
        else:
            raise WordArtifactError(f"不支持的编辑操作：{kind}")
    document.save(output_path)
    return counts


def convert_word_to_pdf(source_path: Path, output_path: Path) -> int:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise WordArtifactError("服务器未安装 LibreOffice，无法转换 PDF")
    with tempfile.TemporaryDirectory(prefix="nano-lo-") as profile:
        environment = {**os.environ, "HOME": profile}
        result = subprocess.run(
            [
                soffice,
                "--headless",
                f"-env:UserInstallation=file://{Path(profile).as_posix()}",
                "--convert-to", "pdf",
                "--outdir", str(output_path.parent),
                str(source_path),
            ],
            capture_output=True,
            text=True,
            timeout=120,
            env=environment,
        )
    generated = output_path.parent / f"{source_path.stem}.pdf"
    if result.returncode != 0 or not generated.exists():
        raise WordArtifactError((result.stderr or result.stdout or "PDF 转换失败")[-1000:])
    if generated != output_path:
        generated.replace(output_path)
    try:
        reader = PdfReader(output_path)
        page_count = len(reader.pages)
    except Exception as error:
        raise WordArtifactError("转换后的 PDF 无法重新打开") from error
    if page_count < 1:
        raise WordArtifactError("转换后的 PDF 没有页面")
    return page_count


def safe_artifact_filename(value: str, extension: str) -> str:
    stem = Path(value).stem.strip() or "nano-document"
    stem = re.sub(r"[\\/:*?\"<>|\x00-\x1f]", "_", stem)[:120].rstrip(". ")
    return f"{stem}{extension}"
